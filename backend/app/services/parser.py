from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

from app.core.config import get_settings
from app.schemas.extraction import TableData

settings = get_settings()


class ParserService:
    def parse(self, file_path: str) -> dict[str, Any]:
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        if ext == ".docx":
            return self._parse_docx(file_path)
        if ext in {".png", ".jpg", ".jpeg", ".webp"}:
            return self._parse_image(file_path)

        raise ValueError(f"Unsupported file type: {ext}")

    # -------------------------
    # PDF PARSING
    # -------------------------
    def _parse_pdf(self, file_path: str) -> dict[str, Any]:
        current_settings = get_settings()
        enable_ai_scan = bool(getattr(current_settings, "ENABLE_AI_SCAN", False))
        ai_provider = str(getattr(current_settings, "AI_PROVIDER", "openai")).lower()

        if enable_ai_scan and ai_provider == "openai":
            try:
                from app.services.ai_scan_parser import AIScanParserService

                ai_parser = AIScanParserService()
                return ai_parser.parse_pdf_to_layout(file_path)
            except Exception as e:
                fallback = self._parse_pdf_without_ai(file_path)
                fallback["warnings"].append(
                    f"AI PDF extraction failed, text fallback used: {str(e)}"
                )
                return fallback

        return self._parse_pdf_without_ai(file_path)

    def _parse_pdf_without_ai(self, file_path: str) -> dict[str, Any]:
        reader = PdfReader(file_path)

        pages_text: list[str] = []
        layout_pages: list[dict[str, Any]] = []
        warnings: list[str] = []

        for page_index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            line_items: list[dict[str, Any]] = []

            if page_text:
                y = 40.0
                for line in page_text.splitlines():
                    line = line.strip()
                    if not line:
                        continue

                    line_items.append(
                        {
                            "type": "text",
                            "bbox": [40.0, y, 760.0, y + 22.0],
                            "text": line,
                        }
                    )
                    y += 26.0
            else:
                warnings.append(
                    f"PDF page {page_index} has no extractable text (OCR disabled)."
                )

            pages_text.append(page_text)

            layout_pages.append(
                {
                    "page_no": page_index,
                    "page_width": 800.0,
                    "page_height": 1000.0,
                    "items": line_items,
                }
            )

        return {
            "raw_text": "\n\n".join([t for t in pages_text if t.strip()]).strip(),
            "tables": [],
            "warnings": warnings,
            "meta": {
                "source_type": "PDF",
                "pages": len(layout_pages),
                "layout_pages": layout_pages,
            },
        }

    # -------------------------
    # DOCX PARSING
    # -------------------------
    def _parse_docx(self, file_path: str) -> dict[str, Any]:
        doc = DocxDocument(file_path)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables: list[TableData] = []

        layout_pages = [
            {
                "page_no": 1,
                "page_width": 800.0,
                "page_height": 1000.0,
                "items": [],
            }
        ]

        current_y = 40.0

        for paragraph in paragraphs:
            layout_pages[0]["items"].append(
                {
                    "type": "text",
                    "bbox": [60.0, current_y, 700.0, current_y + 24.0],
                    "text": paragraph,
                }
            )
            current_y += 32.0

        for idx, table in enumerate(doc.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

            columns = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []

            table_obj = TableData(
                name=f"DOCX_TABLE_{idx}",
                columns=columns,
                rows=data_rows,
            )
            tables.append(table_obj)

            layout_pages[0]["items"].append(
                {
                    "type": "table",
                    "name": table_obj.name,
                    "columns": table_obj.columns,
                    "rows": table_obj.rows,
                    "bbox": [60.0, current_y, 700.0, current_y + 200.0],
                }
            )

            current_y += 240.0

        return {
            "raw_text": "\n".join(paragraphs),
            "tables": tables,
            "warnings": [],
            "meta": {
                "source_type": "DOCX",
                "layout_pages": layout_pages,
            },
        }

    # -------------------------
    # IMAGE PARSING (NO OCR)
    # -------------------------
    def _parse_image(self, file_path: str) -> dict[str, Any]:
        current_settings = get_settings()
        enable_ai_scan = bool(getattr(current_settings, "ENABLE_AI_SCAN", False))
        ai_provider = str(getattr(current_settings, "AI_PROVIDER", "openai")).lower()

        if enable_ai_scan and ai_provider == "openai":
            try:
                from app.services.ai_scan_parser import AIScanParserService

                ai_parser = AIScanParserService()
                return ai_parser.parse_image_to_layout(file_path)
            except Exception as e:
                return self._image_fallback(file_path, f"AI scan failed: {str(e)}")

        return self._image_fallback(file_path)

    def _image_fallback(self, file_path: str, extra_warning: str | None = None):
        with Image.open(file_path) as img:
            width, height = img.size

        warnings = [
            "OCR disabled. Use AI scan or upload PDF/DOCX with text."
        ]

        if extra_warning:
            warnings.insert(0, extra_warning)

        return {
            "raw_text": "",
            "tables": [],
            "warnings": warnings,
            "meta": {
                "source_type": "IMAGE",
                "layout_pages": [
                    {
                        "page_no": 1,
                        "page_width": float(width),
                        "page_height": float(height),
                        "items": [],
                    }
                ],
            },
        }